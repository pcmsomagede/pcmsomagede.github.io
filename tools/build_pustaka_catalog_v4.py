#!/usr/bin/env python3
import json,re,os,time,hashlib,sys
from pathlib import Path
from urllib.parse import urljoin,urlparse
import requests
from bs4 import BeautifulSoup
import fitz
from docx import Document

ROOT=Path('.')
OUT=ROOT/'data/pustaka-catalog.json'
PDF=ROOT/'assets/pustaka-files'; DOCX=ROOT/'assets/pustaka-docx'; COV=ROOT/'assets/pustaka-covers'
PER=100
MAX=int(os.getenv('PUSTAKA_MAX_ITEMS','1000000'))
S=requests.Session(); S.headers.update({'User-Agent':'PCM-Somagede-Pustaka/12.0 (+https://pcmsomagede.github.io/)'})

NOISE=re.compile(r'computer|programming|python|astronomy|physics|chemistry|mathematics|business|marketing|finance|engineering|medicine|biology|travel|cooking|recipe|novel|romance|crime|fiction|sport|football|basketball|patrice|lucas|widner|rogers|monge|washington',re.I)
KH=re.compile(r'khutbah|khutba|jumat|jum[\s\'’]?at|jumah|idul\s+(fitri|adha)|gerhana|istisqa',re.I)
KU=re.compile(r'kultum|kuliah\s+tujuh\s+menit|7\s*menit|ceramah\s+singkat|ceramah\s+pendek|narasi\s+dakwah\s+pendek|pidato\s+agama',re.I)
TF=re.compile(r'tafsir|ta[\'’]wil|tanzil|qur.?an\s+commentary',re.I)
HD=re.compile(r'hadits?|hadith|shahih\s+bukhari|shahih\s+muslim|sahih\s+bukhari|sahih\s+muslim|sunan\s+abu|tirmidzi|tirmidhi|nasai|nasa[\'’]?i|ibnu\s+majah|ibn\s+majah|riyad|riyadh|bulugh',re.I)
ISLAM=re.compile(r'islam|muslim|qur.?an|hadis|hadith|tafsir|fiqih|fikih|aqidah|akidah|tauhid|sunnah|syariah|dakwah|khutbah|jumat|jum.?at|kultum|ceramah|pengajian|pesantren|kitab|sholat|shalat|zakat|puasa|haji|umrah|ramadhan|ramadan|ulama|ustadz|ustaz|doa|ibadah|akhlaq|akhlak|tasawuf|nahwu|sharaf|ushul fiqh|tarjih|muhammadiyah|aisyiyah',re.I)
MUH=re.compile(r'muhammadiyah|tarjih|persyarikatan|suara muhammadiyah|aisyiyah|pemuda muhammadiyah',re.I)

SOURCE_SEEDS={
 'khutbah':[
  ('muhammadiyah','https://muhammadiyah.or.id/?s=khutbah'),
  ('muhammadiyah','https://muhammadiyah.or.id/?s=khutbah+jumat'),
  ('suara-muhammadiyah','https://www.suaramuhammadiyah.id/topic/khutbah'),
  ('tarjih','https://tarjih.or.id/?s=khutbah')],
 'kultum':[
  ('muhammadiyah','https://muhammadiyah.or.id/?s=kultum'),
  ('muhammadiyah','https://muhammadiyah.or.id/?s=ceramah+singkat'),
  ('suara-muhammadiyah','https://www.suaramuhammadiyah.id/?s=kultum'),
  ('tarjih','https://tarjih.or.id/?s=kultum')],
 'tafsir':[
  ('muhammadiyah','https://muhammadiyah.or.id/?s=tafsir'),
  ('tarjih','https://tarjih.or.id/?s=tafsir'),
  ('tarjih-ensiklopedia','https://tarjih.muhammadiyah.or.id/?s=tafsir'),
  ('khazanah','https://khazanah.muhammadiyah.or.id/')],
 'buku':[
  ('tarjih','https://tarjih.or.id/category/gallery/download-file/'),
  ('tarjih','https://tarjih.or.id/?s=Himpunan+Putusan+Tarjih'),
  ('tarjih','https://tarjih.or.id/?s=Tanya+Jawab+Agama'),
  ('khazanah','https://khazanah.muhammadiyah.or.id/')]
}

KEYWORDS={
 'khutbah':['khutbah','khutbah jumat','khutbah idul fitri','khutbah idul adha','khutbah gerhana'],
 'kultum':['kultum','ceramah singkat','kuliah tujuh menit','pidato agama','narasi dakwah'],
 'tafsir':['tafsir','tafsir al quran','tafsir al-quran','tafsir at-tanwir'],
 'buku':['himpunan putusan tarjih','hpt muhammadiyah','tanya jawab agama muhammadiyah','pedoman muhammadiyah','tuntunan muhammadiyah','kitab pesantren','kitab kuning','kitab gundul']
}

def clean(s): return re.sub(r'\s+',' ',str(s or '')).strip(' -|')
def slug(s): return (re.sub(r'[^a-z0-9]+','-',clean(s).lower()).strip('-') or hashlib.sha1(clean(s).encode()).hexdigest()[:16])[:120]
def year(s):
 m=re.search(r'(19|20)\d{2}',str(s or '')); return int(m.group()) if m else 0

def classify(title,section):
 t=clean(title)
 if NOISE.search(t) and not ISLAM.search(t): return False
 if section=='khutbah': return bool(KH.search(t))
 if section=='kultum': return bool(KU.search(t)) and not bool(KH.search(t))
 if section=='tafsir': return bool(TF.search(t))
 return bool(ISLAM.search(t) or re.search(r'pedoman|tuntunan|tarjih|muhammadiyah|kitab|pesantren|fikih|fiqih|aqidah|tauhid|dakwah|kader',t,re.I)) and not bool(KH.search(t) or KU.search(t) or TF.search(t) or HD.search(t))

def soup_text(html):
 s=BeautifulSoup(html,'html.parser')
 for z in s.select('script,style,noscript,nav,header,footer,form,svg,aside'): z.decompose()
 root=s.select_one('article') or s.select_one('.entry-content') or s.select_one('.post-content') or s.select_one('main') or s
 out=[]
 for e in root.select('h1,h2,h3,p,li,blockquote'):
  t=clean(e.get_text(' ',strip=True))
  if t and t not in out: out.append(t)
 return '\n\n'.join(out)

def og_image(html,url):
 s=BeautifulSoup(html,'html.parser'); x=s.select_one('meta[property="og:image"]')
 return urljoin(url,x.get('content')) if x and x.get('content') else ''

def to_pdf(text,title,path):
 try:
  if path.exists(): return True
  d=fitz.open(); p=d.new_page(); y=46; w=p.rect.width-72
  p.insert_text((36,y),title[:115],fontsize=16); y+=28
  for para in re.split(r'\n\s*\n',text):
   line=''
   for word in para.split():
    test=(line+' '+word).strip()
    if fitz.get_text_length(test,fontname='helv',fontsize=10)<=w: line=test
    else:
     if y>p.rect.height-46: p=d.new_page(); y=42
     p.insert_text((36,y),line,fontsize=10); y+=15; line=word
   if line:
    if y>p.rect.height-46: p=d.new_page(); y=42
    p.insert_text((36,y),line,fontsize=10); y+=15
   y+=7
  path.parent.mkdir(parents=True,exist_ok=True); d.save(path); d.close(); return True
 except Exception:return False

def to_docx(text,title,path):
 try:
  if path.exists(): return True
  d=Document(); d.add_heading(title,0)
  for p in re.split(r'\n\s*\n',text):
   p=p.strip()
   if p:d.add_paragraph(p)
  path.parent.mkdir(parents=True,exist_ok=True); d.save(path); return True
 except Exception:return False

def cover_pdf(pp,cp):
 try:
  if cp.exists(): return True
  d=fitz.open(pp); pix=d[0].get_pixmap(matrix=fitz.Matrix(1.7,1.7),alpha=False); cp.parent.mkdir(parents=True,exist_ok=True); pix.save(cp); d.close(); return True
 except Exception:return False

def cover_url(url,cp):
 try:
  if cp.exists():return True
  r=S.get(url,timeout=30); r.raise_for_status();
  if 'image' not in r.headers.get('content-type',''):return False
  cp.parent.mkdir(parents=True,exist_ok=True); cp.write_bytes(r.content); return True
 except Exception:return False

def fetch_page(url):
 try:
  r=S.get(url,timeout=45); r.raise_for_status(); return r.text
 except Exception:return None

def crawl(start,section,group,max_depth=60):
 seen=set(); q=[(start,0)]; out=[]
 host=urlparse(start).netloc
 while q and len(seen)<max_depth:
  u,depth=q.pop(0)
  if u in seen: continue
  seen.add(u); html=fetch_page(u)
  if not html: continue
  s=BeautifulSoup(html,'html.parser')
  for a in s.select('a[href]'):
   href=urljoin(u,a.get('href')); label=clean(a.get_text(' ',strip=True)); net=urlparse(href).netloc
   if net!=host: continue
   low=href.lower()
   if low.endswith('.pdf'):
    if (href,label) not in out: out.append((href,label))
    continue
   blob=(label+' '+href).lower()
   if depth+1<10 and any(k in blob for k in KEYWORDS.get(section,[])):
    q.append((href,depth+1))
 return out

def save_pdf(title,url,section,group):
 try:
  r=S.get(url,timeout=120); r.raise_for_status()
  if not r.content.startswith(b'%PDF') and 'pdf' not in r.headers.get('content-type','').lower(): return None
  n=slug(title); pp=PDF/(n+'.pdf'); dp=DOCX/(n+'.docx'); cp=COV/(n+'.jpg'); pp.parent.mkdir(parents=True,exist_ok=True); pp.write_bytes(r.content)
  d=fitz.open(pp); text='\n\n'.join(pg.get_text('text') for pg in d); d.close()
  if not classify(title,section) or len(text)<100:return None
  if not to_docx(text,title,dp) or not cover_pdf(pp,cp):return None
  return {'title':title,'category':section,'language':'id','source_group':group,'source_page':url,'pdf_source':'/'+str(pp).replace('\\','/'),'docx_source':'/'+str(dp).replace('\\','/'),'cover_url':'/'+str(cp).replace('\\','/'),'year':year(title),'type':'pdf'}
 except Exception:return None

def save_article(title,url,section,group):
 html=fetch_page(url)
 if not html:return None
 title=clean(title)
 if not classify(title,section):return None
 text=soup_text(html)
 if len(text)<300:return None
 n=slug(title); pp=PDF/(n+'.pdf'); dp=DOCX/(n+'.docx'); cp=COV/(n+'.jpg')
 if not to_pdf(text,title,pp) or not to_docx(text,title,dp):return None
 img=og_image(html,url); good=cover_url(img,cp) if img else False
 if not good: good=cover_pdf(pp,cp)
 if not good:return None
 s=BeautifulSoup(html,'html.parser'); tm=s.select_one('time'); dt=tm.get('datetime') if tm else ''
 return {'title':title,'category':section,'language':'id','source_group':group,'source_page':url,'pdf_source':'/'+str(pp).replace('\\','/'),'docx_source':'/'+str(dp).replace('\\','/'),'cover_url':'/'+str(cp).replace('\\','/'),'date':dt or '','year':year(dt or title),'type':'article-export'}

def source_priority(group,section,title=''):
 t=clean(title)
 if section=='khutbah':
  return {'muhammadiyah':700,'suara-muhammadiyah':650,'tarjih':620,'salafi':500,'dewan-dakwah':450,'nu':350}.get(group,100)
 if section=='buku':
  if MUH.search(t): return 800
  return {'muhammadiyah':700,'tarjih':680,'khazanah':650,'pesantren':500,'salafi':400,'nu':300}.get(group,100)
 if section in ('kultum','tafsir'):
  return {'muhammadiyah':700,'suara-muhammadiyah':650,'tarjih':620,'khazanah':600,'kemenag':520,'salafi':450,'pesantren':400,'nu':300}.get(group,100)
 return 100

def key(r): return re.sub(r'\W+','',clean(r.get('title','')).casefold())
def rank(r):
 lang=0 if str(r.get('language','')).lower() in ('id','indonesia','indonesian') else 1
 return (lang,-source_priority(r.get('source_group',''),r.get('category',''),r.get('title','')),-year(r.get('date') or r.get('year')),clean(r.get('title','')).casefold())

def build():
 rows=[]
 # Existing fully-local catalog records are preserved and re-ranked after discovery.
 if OUT.exists():
  try: rows=json.loads(OUT.read_text(encoding='utf-8')).get('items',[])[:]
  except Exception: rows=[]
 # Official Muhammadiyah/SM category/search crawling.
 for section,seeds in SOURCE_SEEDS.items():
  for group,start in seeds:
   urls=crawl(start,section,group,max_depth=140)
   for u,label in urls:
    title=label or Path(urlparse(u).path).stem.replace('-',' ').title()
    if not classify(title,section): continue
    r=save_pdf(title,u,section,group)
    if r: rows.append(r)
   # Follow a limited number of article pages from each seed.
   html=fetch_page(start)
   if not html: continue
   s=BeautifulSoup(html,'html.parser')
   links=[]
   for a in s.select('a[href]'):
    href=urljoin(start,a.get('href')); label=clean(a.get_text(' ',strip=True))
    if urlparse(href).netloc==urlparse(start).netloc and classify(label,section) and href not in links: links.append(href)
   for href in links[:120]:
    r=save_article(clean(BeautifulSoup(fetch_page(href) or '', 'html.parser').title.get_text(' ',strip=True) if BeautifulSoup(fetch_page(href) or '', 'html.parser').title else ''),href,section,group)
    if r: rows.append(r)
 # Curated official public PDFs known to be real.
 known=[
  ('buku','Hukum Takziah dan Ziarah Kubur','https://tarjih.or.id/wp-content/uploads/2022/01/Hukum-Takziah-dan-Ziarah-Kubur.pdf','tarjih'),
  ('buku','Tuntunan Walimah','https://tarjih.or.id/wp-content/uploads/2020/08/ebook-Tuntunan-Walimah.pdf','tarjih'),
  ('buku','Kajian Tarjih Fidyah dan Zakat Fitri','https://tarjih.or.id/wp-content/uploads/2021/05/Kajian-Tarjih-Fidyah-Zakat-FITRI-Idul-Fitri-Ali.pdf','tarjih')]
 for sec,t,u,g in known:
  r=save_pdf(t,u,sec,g)
  if r: rows.append(r)
 # Strict final gate: all visible downloadable/preview assets are local.
 clean_rows=[]; seen=set()
 for r in rows:
  if not classify(r.get('title',''),r.get('category','')):continue
  paths=[r.get('pdf_source',''),r.get('docx_source',''),r.get('cover_url','')]
  if not all(str(x).startswith('/assets/') for x in paths):continue
  k=key(r)
  if not k or k in seen:continue
  seen.add(k); clean_rows.append(r)
 clean_rows.sort(key=rank)
 OUT.parent.mkdir(parents=True,exist_ok=True)
 OUT.write_text(json.dumps({'generated':time.strftime('%Y-%m-%dT%H:%M:%SZ',time.gmtime()),'total':len(clean_rows[:MAX]),'items':clean_rows[:MAX]},ensure_ascii=False,indent=2)+'\n',encoding='utf-8')
 print('PustakaMu usable records:',len(clean_rows[:MAX]))

if __name__=='__main__': build()
