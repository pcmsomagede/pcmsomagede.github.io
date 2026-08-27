#!/usr/bin/env python3
import json,re,time,os,io,hashlib
from pathlib import Path
from urllib.parse import urljoin,urlparse
import requests
from bs4 import BeautifulSoup
try:
 import fitz
except Exception:
 fitz=None
try:
 from docx import Document
except Exception:
 Document=None

OUT=Path('data/pustaka-catalog.json')
PDF_DIR=Path('assets/pustaka-files'); COVER_DIR=Path('assets/pustaka-covers'); DOCX_DIR=Path('assets/pustaka-docx')
MAX_ITEMS=int(os.getenv('PUSTAKA_MAX_ITEMS','1000000'))
PER_PAGE=100
HEADERS={'User-Agent':'PCM-Somagede-Pustaka/7.0 (+https://pcmsomagede.github.io/)'}
S=requests.Session();S.headers.update(HEADERS)

SOURCES={
 'muhammadiyah':['https://muhammadiyah.or.id','https://tarjih.or.id','https://suaramuhammadiyah.id','https://buku.muhammadiyah.or.id'],
 'kemenag':['https://kemenag.go.id','https://quran.kemenag.go.id'],
 'salafi':['https://yufid.com','https://rodja.com','https://rodja.tv','https://muslimafiyah.com','https://raehanulbahrain.com'],
 'nu':['https://nu.or.id']
}
ALLOWED_HOSTS=tuple(sorted({urlparse(x).hostname for xs in SOURCES.values() for x in xs if urlparse(x).hostname}))
NOISE=re.compile(r'computer|programming|python|astronomy|physics|chemistry|mathematics|business|marketing|finance|engineering|medicine|biology|travel|cooking|recipe|novel|romance|crime|fiction|patrice|lucas|widner|rogers|monge|washington',re.I)
KH=re.compile(r'khutbah|khutba|jumat|jum\\s*[\'’]?at|jumah|idul\\s+(fitri|adha)|gerhana|istisqa',re.I)
KU=re.compile(r'kultum|kuliah\\s+tujuh\\s+menit|7\\s*menit|ceramah\\s+(singkat|pendek)|narasi\\s+dakwah\\s+pendek|pidato\\s+agama',re.I)
TF=re.compile(r'tafsir|tafs?ir|tanzil|ta.wil|qur.?an\\s+commentary',re.I)
HD=re.compile(r'hadits?|hadith|shahih\\s+bukhari|shahih\\s+muslim|sunan\\s+abu|abu\\s+dawud|tirmidzi|tirmidhi|nasai|nasa.i|ibnu\\s+majah|ibn\\s+majah|riyad|bulugh',re.I)
ISLAM=re.compile(r'islam|muslim|qur.?an|hadis|hadith|tafsir|fiqih|fikih|aqidah|akidah|tauhid|sunnah|syariah|dakwah|khutbah|jumat|kultum|ceramah|pengajian|pesantren|kitab|sholat|shalat|zakat|puasa|haji|umrah|ramadhan|ulama|ustadz|doa|ibadah|akhlak|tasawuf|nahwu|sharaf|ushul\\s+fiqh|tarjih|muhammadiyah',re.I)

SEED_PAGES=[
 ('khutbah','https://muhammadiyah.or.id/2026/08/khutbah-jumat-kemerdekaan-dalam-bingkai-baldatun-thayyibatun-wa-rabbun-ghafur/'),
 ('khutbah','https://muhammadiyah.or.id/2026/07/khutbah-jumat-taaruf-sebagai-jalan-menuju-keluarga-sakinah/'),
 ('khutbah','https://muhammadiyah.or.id/2026/06/khutbah-jumat-menyelaraskan-qalb-aql-dan-nafs-menuju-jiwa-yang-tenang/'),
 ('khutbah','https://muhammadiyah.or.id/2026/06/khutbah-jumat-menempatkan-amanah-jabatan-pada-ahlinya/'),
 ('khutbah','https://muhammadiyah.or.id/2026/06/khutbah-jumat-keutamaan-puasa-tasua-dan-asyura/'),
 ('khutbah','https://muhammadiyah.or.id/2026/05/khutbah-jumat-bahaya-rokok-dalam-perspektif-al-quran-dan-sunnah/'),
 ('khutbah','https://muhammadiyah.or.id/2026/01/khutbah-jumat-menjemput-ramadan-dengan-puasa-sunah-di-bulan-syaban/'),
 ('khutbah','https://muhammadiyah.or.id/2025/08/khutbah-jumat-mewujudkan-negara-pancasila-sebagai-darul-ahdi-wa-syahadah/'),
 ('kultum','https://muhammadiyah.or.id/2026/02/kultum-ramadan-ukhuwah-di-bulan-yang-suci/'),
 ('kultum','https://muhammadiyah.or.id/2026/02/kultum-ramadan-memaknai-amal-jariyah-lebih-luas/'),
 ('kultum','https://muhammadiyah.or.id/2026/03/kultum-ramadan-masjid-sebagai-pusat-kemanusiaan/'),
 ('kultum','https://muhammadiyah.or.id/2026/02/hukum-kultum-sebelum-salat-tarawih/'),
 ('kultum','https://muhammadiyah.or.id/2024/02/agar-kultum-ramadhan-tidak-terasa-membosankan/'),
]

def host_ok(u):
 h=urlparse(u).hostname or ''
 return any(h==x or h.endswith('.'+x) for x in ALLOWED_HOSTS)

def txt(v):
 if isinstance(v,list): return ' '.join(map(str,v))
 return str(v or '')

def normalize(s): return re.sub(r'[^a-z0-9]+','',txt(s).lower())

def classify(t):
 if KH.search(t): return 'khutbah'
 if KU.search(t): return 'kultum'
 if TF.search(t): return 'tafsir'
 if HD.search(t): return 'hadits'
 return 'buku'

def lang(v): return 'id' if re.search(r'indonesia|indonesian|bahasa indonesia|\\bid\\b',txt(v).lower()) else 'other'

def year(v):
 m=re.search(r'(19|20)\\d{2}',txt(v)); return int(m.group()) if m else 0

def clean(base,u): return urljoin(base,txt(u).split('#')[0].strip()) if u else ''

def safe_name(t,ext):
 s=re.sub(r'[^a-z0-9]+','-',txt(t).lower()).strip('-')[:80] or 'dokumen'
 return s+ext

def write_pdf_from_text(text,title,out):
 if not fitz: return False
 try:
  doc=fitz.open(); page=doc.new_page(); w=page.rect.width-72; y=55
  page.insert_text((36,y),title[:110],fontsize=16,fontname='helv'); y+=28
  for para in re.split(r'\\n\\s*\\n',text):
   words=para.split(); line=''
   for word in words:
    test=(line+' '+word).strip()
    if fitz.get_text_length(test,fontname='helv',fontsize=10)<=w: line=test
    else:
     if line:
      if y>page.rect.height-45: page=doc.new_page(); y=45
      page.insert_text((36,y),line,fontsize=10,fontname='helv'); y+=15
     line=word
   if line:
    if y>page.rect.height-45: page=doc.new_page(); y=45
    page.insert_text((36,y),line,fontsize=10,fontname='helv'); y+=15
   y+=8
  out.parent.mkdir(parents=True,exist_ok=True);doc.save(out);doc.close();return True
 except Exception:return False

def write_docx(text,title,out):
 if not Document:return False
 try:
  out.parent.mkdir(parents=True,exist_ok=True);doc=Document();doc.add_heading(title,0)
  for p in re.split(r'\\n\\s*\\n',text):
   p=p.strip()
   if p: doc.add_paragraph(p)
  doc.save(out);return True
 except Exception:return False

def article_text(html):
 s=BeautifulSoup(html,'html.parser')
 for x in s.select('script,style,noscript,nav,header,footer,form'):
  x.decompose()
 root=s.select_one('article') or s.select_one('.entry-content') or s.select_one('.post-content') or s.select_one('main') or s
 title=root.select_one('h1') if root else None
 blocks=[]
 for el in root.select('h1,h2,h3,p,li'):
  v=el.get_text(' ',strip=True)
  if v and v not in blocks: blocks.append(v)
 return '\\n\\n'.join(blocks)

def seed_articles():
 out=[]
 for category,url in SEED_PAGES:
  try:
   r=S.get(url,timeout=45); r.raise_for_status(); s=BeautifulSoup(r.text,'html.parser')
   h=s.select_one('h1') or s.title; title=h.get_text(' ',strip=True) if h else url.rsplit('/',2)[-2]
   body=article_text(r.text)
   if len(body)<300: continue
   pdfp=PDF_DIR/safe_name(title,'.pdf'); docxp=DOCX_DIR/safe_name(title,'.docx'); coverp=COVER_DIR/safe_name(title,'.jpg')
   okpdf=pdfp.exists() or write_pdf_from_text(body,title,pdfp)
   okdoc=docxp.exists() or write_docx(body,title,docxp)
   cov=''
   if okpdf and fitz and (not coverp.exists()):
    try:
     d=fitz.open(pdfp); pix=d[0].get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False);coverp.parent.mkdir(parents=True,exist_ok=True);pix.save(coverp);d.close()
    except Exception: pass
   if coverp.exists(): cov='/assets/pustaka-covers/'+coverp.name
   if okpdf or okdoc:
    out.append({'title':title,'category':category,'source_page':url,'date':re.search(r'/20\\d{2}/(?:0[1-9]|1[0-2])/',url).group(0).strip('/') if re.search(r'/20\\d{2}/(?:0[1-9]|1[0-2])/',url) else '','pdf_source':'/'+str(pdfp).replace('\\\\','/'),'docx_source':'/'+str(docxp).replace('\\\\','/'),'cover_url':cov,'source_group':'muhammadiyah','language':'id','type':'article-export'})
  except Exception: pass
 return out

def direct_pdfs():
 roots=['https://tarjih.or.id/category/gallery/download-file/','https://tarjih.or.id/category/produk/putusan/','https://tarjih.or.id/category/produk/wacana/','https://tarjih.or.id/category/produk/fatwa/','https://muhammadiyah.or.id/download/','https://muhammadiyah.or.id/category/publikasi/']
 out=[]
 for root in roots:
  try:
   s=BeautifulSoup(S.get(root,timeout=45).text,'html.parser')
   for a in s.select('a[href]'):
    u=clean(root,a.get('href'))
    if host_ok(u) and re.search(r'\\.pdf(?:$|\\?)',urlparse(u).path,re.I):
     t=a.get_text(' ',strip=True) or re.sub(r'\\.pdf$','',u.rsplit('/',1)[-1],flags=re.I)
     out.append({'title':t,'category':classify(t),'source_page':root,'remote_pdf':u,'source_group':'muhammadiyah','language':'id'})
  except Exception: pass
 return out

def localize_pdf(r):
 u=r.get('remote_pdf') or r.get('pdf_source')
 if not u:return r
 if u.startswith('/'):
  p=Path(u.lstrip('/')); r['pdf_source']=u; return r if p.exists() else {}
 try:
  resp=S.get(u,timeout=90);resp.raise_for_status();name=safe_name(r['title'],'.pdf');p=PDF_DIR/name;p.parent.mkdir(parents=True,exist_ok=True);p.write_bytes(resp.content);r['pdf_source']='/'+str(p).replace('\\\\','/')
  if fitz:
   cp=COVER_DIR/(Path(name).stem+'.jpg')
   if not cp.exists():
    d=fitz.open(stream=resp.content,filetype='pdf');pix=d[0].get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False);cp.parent.mkdir(parents=True,exist_ok=True);pix.save(cp);d.close()
   r['cover_url']='/assets/pustaka-covers/'+cp.name
   dp=DOCX_DIR/(Path(name).stem+'.docx');
   if not dp.exists() and Document:
    d=fitz.open(stream=resp.content,filetype='pdf');text='\\n\\n'.join(pg.get_text('text') for pg in d);d.close();write_docx(text,r['title'],dp)
   if dp.exists():r['docx_source']='/assets/pustaka-docx/'+dp.name
  return r
 except Exception:return {}

def wp_posts(base,search,pages=4):
 out=[]
 for p in range(1,pages+1):
  try:
   r=S.get(base+'/wp-json/wp/v2/posts',params={'search':search,'per_page':PER_PAGE,'page':p,'_fields':'link,date,modified,title,content,excerpt'},timeout=45)
   if not r.ok:break
   rows=r.json();out.extend(rows)
   if len(rows)<PER_PAGE:break
  except Exception:break
 return out

def post_record(group,p):
 t=BeautifulSoup(txt(p.get('title',{}).get('rendered') if isinstance(p.get('title'),dict) else p.get('title')),'html.parser').get_text(' ',strip=True)
 html=txt(p.get('content',{}).get('rendered') if isinstance(p.get('content'),dict) else p.get('content'));link=txt(p.get('link'));c=classify(t+' '+html+' '+link);return {'title':t,'category':c,'source_page':link,'date':p.get('date') or p.get('modified') or '','source_group':group,'language':'id'}

def crawl():
 rec=[];terms={'muhammadiyah':['Himpunan Putusan Tarjih','Tanya Jawab Agama','pedoman','tarjih','dakwah','tafsir','khutbah','kultum','ibadah','keislaman'],'kemenag':['tafsir','quran','Islam','khutbah','ceramah'],'salafi':['tafsir','khutbah','kultum','ceramah','kitab','Islam'],'nu':['khutbah','tafsir','ceramah','kitab','Islam']}
 for group,bases in SOURCES.items():
  for base in bases:
   if group=='muhammadiyah' and 'buku.' in base: continue
   for term in terms[group]:
    for p in wp_posts(base,term,pages=2):
     try:rec.append(post_record(group,p))
     except Exception:pass
 return rec

def merge(rows):
 old={normalize(x.get('title')):x for x in json.loads(OUT.read_text(encoding='utf-8')).get('items',[])} if OUT.exists() else {}
 merged={}
 for r in rows:
  title=txt(r.get('title')).strip()
  if not title:continue
  blob=txt(r)
  if NOISE.search(title) and not ISLAM.search(blob):continue
  if r.get('category')=='hadits': continue
  k=normalize(title);prev=merged.get(k) or old.get(k);rr=dict(prev or {});rr.update({a:b for a,b in r.items() if b});merged[k]=rr
 # localize remote docs, and only keep records with local readable assets
 out=[]
 for r in merged.values():
  if r.get('remote_pdf'):
   r=localize_pdf(r)
  p=r.get('pdf_source','');d=r.get('docx_source','');c=r.get('cover_url','')
  if p and not p.startswith('/'):p='/' + p
  if d and not d.startswith('/'):d='/' + d
  if c and not c.startswith('/'):c='/' + c
  r['pdf_source']=p;r['docx_source']=d;r['cover_url']=c
  if not (p or d): continue
  # never publish a broken local asset
  if p and not Path(p.lstrip('/')).exists(): p='';r['pdf_source']=''
  if d and not Path(d.lstrip('/')).exists(): d='';r['docx_source']=''
  if not p and not d:continue
  if not c and p and fitz:
   try:
    pp=Path(p.lstrip('/'));cp=COVER_DIR/(pp.stem+'.jpg');cp.parent.mkdir(parents=True,exist_ok=True)
    if not cp.exists():
     dd=fitz.open(pp);pix=dd[0].get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False);pix.save(cp);dd.close()
    r['cover_url']='/assets/pustaka-covers/'+cp.name
   except Exception:pass
  out.append(r)
 def sk(r):
  g=r.get('source_group','lain');gs={'muhammadiyah':1000,'kemenag':850,'salafi':700,'nu':500,'lain':100}.get(g,100);cat={'khutbah':4,'kultum':3,'buku':2,'tafsir':1}.get(r.get('category'),0)
  return (0 if lang(r.get('language'))=='id' else 1,-gs,-cat,-year(r.get('date') or r.get('year')),txt(r.get('title')).lower())
 out.sort(key=sk);return out[:MAX_ITEMS]

def build():
 rows=seed_articles()+direct_pdfs()+crawl();items=merge(rows)
 OUT.parent.mkdir(parents=True,exist_ok=True);json.dump({'generated':time.strftime('%Y-%m-%dT%H:%M:%SZ',time.gmtime()),'total':len(items),'items':items},OUT.open('w',encoding='utf-8'),ensure_ascii=False,indent=2)
 print('PustakaMu local records:',len(items))
if __name__=='__main__':build()
