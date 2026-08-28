#!/usr/bin/env python3
"""Regenerate HTML-exported khutbah PDFs/DOCX with Arabic/RTL preserved."""
from __future__ import annotations
import json,re
from pathlib import Path
from urllib.parse import urljoin
import requests
from bs4 import BeautifulSoup
from docx import Document
from weasyprint import HTML

CAT=Path('data/pustaka-catalog.json')
PDF=Path('assets/pustaka-files'); DOCX=Path('assets/pustaka-docx'); COV=Path('assets/pustaka-covers')
AR=re.compile(r'[\u0600-\u06ff\u0750-\u077f\u08a0-\u08ff]')
S=requests.Session(); S.headers.update({'User-Agent':'PCM-Somagede-KhutbahRepair/1.0'})

def clean(v): return re.sub(r'\s+',' ',str(v or '')).strip()
def slug(v): return re.sub(r'[^a-z0-9]+','-',v.lower()).strip('-')[:110]
def canonical_title(v): return re.sub(r'\s*\|\s*Muhammadiyah\s*$','',clean(v),flags=re.I).casefold()

def sanitize(html,title):
    s=BeautifulSoup(html,'html.parser')
    for z in s.select('script,style,noscript,nav,header,footer,form,svg,iframe,.share,.sharing,.comments,.comment,.related,.advertisement,.ads,.sidebar'):
        z.decompose()
    root=s.select_one('article') or s.select_one('.entry-content') or s.select_one('.post-content') or s.select_one('main') or s.body or s
    keep=[]
    for e in root.find_all(['h1','h2','h3','h4','p','ol','ul','li','blockquote','strong','em','br','table','tr','td','th']):
        t=e.get_text(' ',strip=True)
        if not t: continue
        if AR.search(t):
            e['dir']='rtl'; e['class']=list(e.get('class',[]))+['arabic-present']
        keep.append(e)
    body=''.join(str(e) for e in keep)
    return f'''<!doctype html><html lang="id"><head><meta charset="utf-8"><style>
@page{{size:A4;margin:16mm 15mm}}
body{{font-family:"Noto Sans","Noto Sans Arabic",sans-serif;font-size:11pt;line-height:1.7;color:#161616}}
h1,h2,h3,h4{{font-family:"Noto Sans","Noto Sans Arabic",sans-serif;line-height:1.3}}
p,li,blockquote{{margin:.45em 0}}
.arabic-present{{direction:rtl;unicode-bidi:plaintext;text-align:right;font-family:"Noto Naskh Arabic","Noto Sans Arabic","Amiri Quran",serif;font-size:15pt;line-height:2.05}}
.arabic-present *{{font-family:inherit}}
blockquote{{border-left:3px solid #aaa;padding-left:10px}}
table{{border-collapse:collapse;width:100%}}td,th{{border:1px solid #aaa;padding:4px}}
</style></head><body><h1>{title}</h1>{body}</body></html>'''

def regen(item):
    url=item.get('source_page') or ''
    if not url or re.search(r'\.pdf(?:$|\?)',url,re.I): return False
    r=S.get(url,timeout=90); r.raise_for_status()
    html=sanitize(r.text,clean(item.get('title') or 'Khutbah'))
    title=clean(item.get('title') or '')
    base=slug(canonical_title(title)) or slug(title)
    pp=PDF/(base+'.pdf'); dp=DOCX/(base+'.docx'); cp=COV/(base+'.jpg')
    pp.parent.mkdir(parents=True,exist_ok=True); dp.parent.mkdir(parents=True,exist_ok=True)
    HTML(string=html,base_url=url).write_pdf(pp)
    soup=BeautifulSoup(html,'html.parser'); d=Document(); d.add_heading(title,0)
    for e in soup.find_all(['p','li','blockquote','h2','h3','h4']):
        t=e.get_text('\n',strip=True)
        if not t: continue
        p=d.add_paragraph(t)
        if AR.search(t): p.alignment=2
    d.save(dp)
    try:
        import fitz
        doc=fitz.open(pp); pix=doc[0].get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False); cp.parent.mkdir(parents=True,exist_ok=True); pix.save(cp); doc.close()
    except Exception: pass
    # Verify that article-export output really contains Arabic text in the PDF text layer.
    import fitz
    doc=fitz.open(pp); extracted='\n'.join(p.get_text('text') for p in doc); doc.close()
    if not AR.search(extracted):
        raise RuntimeError(f'Arabic verification failed: {title}')
    item['pdf_source']='/'+str(pp).replace('\\','/')
    item['docx_source']='/'+str(dp).replace('\\','/')
    if cp.exists(): item['cover_url']='/'+str(cp).replace('\\','/')
    item['type']='article-export-arabic-safe'
    return True

def main():
    if not CAT.exists(): return
    cat=json.loads(CAT.read_text(encoding='utf-8')); out=[];seen=set(); repaired=0
    for item in cat.get('items') or []:
        catg=str(item.get('category') or '').casefold()
        if catg=='khutbah':
            try:
                if str(item.get('type') or '').startswith('article-export') or item.get('source_page'):
                    if regen(item): repaired+=1
            except Exception as e:
                print('REPAIR FAILED:',item.get('title'),e)
                raise
        key=(catg,canonical_title(item.get('title')),clean(item.get('source_page')).casefold())
        if key in seen: continue
        seen.add(key); out.append(item)
    cat['items']=out; cat['total']=len(out)
    CAT.write_text(json.dumps(cat,ensure_ascii=False,indent=2)+'\n',encoding='utf-8')
    print(f'Khutbah exports repaired: {repaired}; catalog records: {len(out)}')

if __name__=='__main__': main()
