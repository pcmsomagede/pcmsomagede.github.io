#!/usr/bin/env python3
import json,re,os,time,hashlib
from pathlib import Path
from urllib.parse import urljoin,urlparse,quote
import requests
from bs4 import BeautifulSoup
import fitz
from docx import Document

OUT=Path('data/pustaka-catalog.json'); PDF=Path('assets/pustaka-files'); DOCX=Path('assets/pustaka-docx'); COV=Path('assets/pustaka-covers')
MAX=int(os.getenv('PUSTAKA_MAX_ITEMS','1000000')); PER=100
S=requests.Session(); S.headers.update({'User-Agent':'PCM-Somagede-Pustaka/11.0 (+https://pcmsomagede.github.io/)'})
NOISE=re.compile(r'computer|programming|python|astronomy|physics|chemistry|mathematics|business|marketing|finance|engineering|medicine|biology|travel|cooking|recipe|novel|romance|crime|fiction|patrice|lucas|widner|rogers|monge|washington',re.I)
ISLAM=re.compile(r'islam|muslim|qur.?an|hadis|hadith|tafsir|fiqih|fikih|aqidah|akidah|tauhid|sunnah|syariah|dakwah|khutbah|jumat|jum.?at|kultum|ceramah|pengajian|pesantren|kitab|sholat|shalat|zakat|puasa|haji|umrah|ramadhan|ramadan|ulama|ustadz|ustaz|doa|ibadah|akhlaq|akhlak|tasawuf|nahwu|sharaf|ushul fiqh|tarjih|muhammadiyah|aisyiyah',re.I)
KH=re.compile(r'khutbah|khutba|jumat|jum.?at|jumah|idul\s+(fitri|adha)|gerhana|istisqa',re.I)
KU=re.compile(r'kultum|kuliah\s+tujuh\s+menit|7\s*menit|ceramah\s+singkat|ceramah\s+pendek|narasi\s+dakwah\s+pendek|pidato\s+agama',re.I)
TF=re.compile(r'tafsir|tanzil|ta.?wil|qur.?an\s+commentary',re.I)
HD=re.compile(r'hadits?|hadith|shahih\s+bukhari|shahih\s+muslim|sahih\s+bukhari|sahih\s+muslim|sunan\s+abu|tirmidzi|tirmidhi|nasai|nasa.?i|ibnu\s+majah|ibn\s+majah|riyad|bulugh',re.I)

STARTS={
 'khutbah':['https://muhammadiyah.or.id/khutbah-jumat/','https://muhammadiyah.or.id/khutbah-/'],
 'kultum':['https://muhammadiyah.or.id/?s=kultum','https://muhammadiyah.or.id/?s=ceramah+singkat'],
 'tafsir':['https://tarjih.or.id/?s=tafsir','https://muhammadiyah.or.id/?s=tafsir'],
 'buku':['https://tarjih.or.id/category/gallery/download-file/','https://tarjih.or.id/?s=Himpunan+Putusan+Tarjih','https://tarjih.or.id/?s=Tanya+Jawab+Agama']
}
KNOWN=[
 ('buku','Tuntunan Thaharah','https://tarjih.or.id/wp-content/uploads/2020/08/ebook-Tuntunan-Thaharah.pdf','muhammadiyah'),
 ('buku','Tuntunan Walimah','https://tarjih.or.id/wp-content/uploads/2020/08/ebook-Tuntunan-Walimah.pdf','muhammadiyah'),
 ('buku','Hukum Takziah dan Ziarah Kubur','https://tarjih.or.id/wp-content/uploads/2022/01/Hukum-Takziah-dan-Ziarah-Kubur.pdf','muhammadiyah'),
 ('buku','Tajdid Muhammadiyah untuk Pencerahan Peradaban','https://tarjih.or.id/wp-content/uploads/2020/08/TAJDID-MUHAMMADIYAH-UNTUK-PENCERAHAN-PERADABAN.pdf','muhammadiyah'),
 ('buku','Memahami Paham Agama Muhammadiyah','https://tarjih.or.id/wp-content/uploads/2021/05/Makalah-Posisi-Paham-Agama-Muhammadiyah-dari-Manhaj-Tarjih-Terbaru.pdf','muhammadiyah'),
 ('buku','Manhaj Tarjih Muhammadiyah','https://tarjih.or.id/wp-content/uploads/2020/09/Manhaj-Tarjih-Pengajian-Masjid-Gede-Edisi-2-Tahun-2018-a-1.pdf','muhammadiyah'),
 ('buku','Aplikasi Manhaj Tarjih: Tanawwu Fil Ibadah','https://tarjih.or.id/wp-content/uploads/2022/12/KONSEP-AT-TANAWWU-FIL-IBADAH-DAN-CONTOH-CONTOHNYA-DALAM-HPT.pdf','muhammadiyah'),
 ('buku','Kapita Selekta Putusan dan Fatwa Tarjih Bidang Akidah','https://tarjih.or.id/wp-content/uploads/2021/05/KAPITA-SELEKTA-PUTUSAN-DAN-FATWA-TARJIH-BIDANG-AKIDAH.pdf','muhammadiyah'),
 ('tafsir','Tafsir At-Tanwir: Al-Baqarah 130–134','https://tarjih.or.id/wp-content/uploads/2021/03/Tafsir-al-Baqarah-Ayat-130-134-materi-pengajian-tarjih-edisi-120.pdf','muhammadiyah'),
 ('buku','Kajian Tarjih Fidyah dan Zakat Fitri','https://tarjih.or.id/wp-content/uploads/2021/05/Kajian-Tarjih-Fidyah-Zakat-FITRI-Idul-Fitri-Ali.pdf','muhammadiyah'),
 ('buku','Berita Resmi Tanfidz Keputusan Munas Tarjih','https://tarjih.or.id/wp-content/uploads/2024/03/BRM-03-Tanfidz-Keputusan-Munas-Tarjih.pdf','muhammadiyah')
]

def slug(t):
 s=re.sub(r'[^a-z0-9]+','-',t.lower()).strip('-'); return (s[:110] or hashlib.sha1(t.encode()).hexdigest()[:16])
def clean_title(t): return re.sub(r'\s+',' ',str(t or '')).strip(' -|')
def date_year(v):
 m=re.search(r'(19|20)\d{2}',str(v or '')); return int(m.group()) if m else 0

def pdf_from_text(text,title,path):
 try:
  if path.exists(): return True
  d=fitz.open(); pg=d.new_page(); y=46; w=pg.rect.width-72; pg.insert_text((36,y),title[:110],fontsize=16); y+=28
  for para in re.split(r'\n\s*\n',text):
   line=''
   for word in para.split():
    test=(line+' '+word).strip()
    if fitz.get_text_length(test,fontname='helv',fontsize=10)<=w: line=test
    else:
     if y>pg.rect.height-46: pg=d.new_page(); y=42
     pg.insert_text((36,y),line,fontsize=10); y+=15; line=word
   if line:
    if y>pg.rect.height-46: pg=d.new_page(); y=42
    pg.insert_text((36,y),line,fontsize=10); y+=15
   y+=7
  path.parent.mkdir(parents=True,exist_ok=True); d.save(path); d.close(); return True
 except Exception:return False

def docx_from_text(text,title,path):
 try:
  if path.exists(): return True
  path.parent.mkdir(parents=True,exist_ok=True); d=Document(); d.add_heading(title,0)
  for p in re.split(r'\n\s*\n',text):
   p=p.strip()
   if p:d.add_paragraph(p)
  d.save(path); return True
 except Exception:return False

def cover_from_pdf(pdf_path,path):
 try:
  if path.exists(): return True
  d=fitz.open(pdf_path); pix=d[0].get_pixmap(matrix=fitz.Matrix(1.6,1.6),alpha=False); path.parent.mkdir(parents=True,exist_ok=True); pix.save(path); d.close(); return True
 except Exception:return False

def cover_from_image(url,path):
 try:
  if path.exists(): return True
  r=S.get(url,timeout=30); r.raise_for_status(); ct=r.headers.get('content-type','')
  if 'image' not in ct:return False
  path.parent.mkdir(parents=True,exist_ok=True); path.write_bytes(r.content); return True
 except Exception:return False

def html_text(html):
 s=BeautifulSoup(html,'html.parser')
 for z in s.select('script,style,noscript,nav,header,footer,form,svg'): z.decompose()
 root=s.select_one('article') or s.select_one('.entry-content') or s.select_one('.post-content') or s.select_one('main') or s
 parts=[]
 for e in root.select('h1,h2,h3,p,li,blockquote'):
  t=e.get_text(' ',strip=True)
  if t and t not in parts: parts.append(t)
 return '\n\n'.join(parts)

def og_image(html,url):
 s=BeautifulSoup(html,'html.parser'); x=s.select_one('meta[property="og:image"]'); return urljoin(url,x.get('content')) if x and x.get('content') else ''

def in_section(title,section):
 t=clean_title(title)
 if NOISE.search(t) and not ISLAM.search(t): return False
 if section=='khutbah': return bool(KH.search(t))
 if section=='kultum': return bool(KU.search(t)) and not bool(KH.search(t))
 if section=='tafsir': return bool(TF.search(t))
 if section=='buku': return not bool(KH.search(t) or KU.search(t) or TF.search(t) or HD.search(t)) and bool(ISLAM.search(t) or re.search(r'pedoman|tuntunan|tarjih|muhammadiyah|kitab|fikih|fiqih|akidah|tauhid|dakwah|kader',t,re.I))
 return False

def article_record(section,url,group='muhammadiyah'):
 try:
  r=S.get(url,timeout=45); r.raise_for_status(); html=r.text; s=BeautifulSoup(html,'html.parser'); h=s.select_one('h1') or s.title; title=clean_title(h.get_text(' ',strip=True) if h else url)
  if not in_section(title,section): return None
  text=html_text(html)
  if len(text)<260:return None
  n=slug(title); pp=PDF/(n+'.pdf'); dp=DOCX/(n+'.docx'); cp=COV/(n+'.jpg')
  if not pdf_from_text(text,title,pp):return None
  if not docx_from_text(text,title,dp):return None
  img=og_image(html,url); ok=cover_from_image(img,cp) if img else False
  if not ok: ok=cover_from_pdf(pp,cp)
  if not ok:return None
  meta={'title':title,'category':section,'source_page':url,'language':'id','source_group':group,'date':s.select_one('time') and s.select_one('time').get('datetime') or '', 'pdf_source':'/'+str(pp).replace('\\','/'),'docx_source':'/'+str(dp).replace('\\','/'),'cover_url':'/'+str(cp).replace('\\','/'),'type':'article-export'}
  return meta
 except Exception:return None

def discover_html(start,section,max_pages=40):
 out=[]; seen=set(); queue=[start]
 for _ in range(max_pages):
  if not queue:break
  u=queue.pop(0)
  if u in seen:continue
  seen.add(u)
  try:
   r=S.get(u,timeout=35); r.raise_for_status(); s=BeautifulSoup(r.text,'html.parser')
   for a in s.select('a[href]'):
    href=urljoin(u,a.get('href')); txt=clean_title(a.get_text(' ',strip=True));
    if urlparse(href).netloc not in {'muhammadiyah.or.id','tarjih.or.id'}: continue
    low=href.lower()
    if low.endswith('.pdf'):
     if href not in [x[1] for x in out]: out.append(('pdf',href,txt))
    elif any(k in (txt+' '+href).lower() for k in ['khutbah','kultum','tafsir','tarjih','himpunan putusan','tanya jawab agama','tuntunan','pedoman']):
     if href not in seen and len(queue)<200: queue.append(href)
   nxt=s.select_one('a[rel="next"]') or s.select_one('a.next')
   if nxt and nxt.get('href'):queue.append(urljoin(u,nxt.get('href')))
  except Exception:continue
 return out

def wp_posts(base,term,max_pages=30):
 out=[]
 for p in range(1,max_pages+1):
  try:
   r=S.get(base+'/wp-json/wp/v2/posts',params={'search':term,'per_page':PER,'page':p,'_fields':'link,date,modified,title,content'},timeout=40)
   if not r.ok:break
   rows=r.json(); out+=rows
   if len(rows)<PER:break
  except Exception:break
 return out

def remote_pdf(title,url,cat,group='muhammadiyah'):
 try:
  r=S.get(url,timeout=90); r.raise_for_status()
  if 'pdf' not in r.headers.get('content-type','').lower() and not r.content.startswith(b'%PDF'): return None
  n=slug(title); pp=PDF/(n+'.pdf'); dp=DOCX/(n+'.docx'); cp=COV/(n+'.jpg'); pp.parent.mkdir(parents=True,exist_ok=True); pp.write_bytes(r.content)
  d=fitz.open(pp); text='\n\n'.join(pg.get_text('text') for pg in d); d.close()
  if not docx_from_text(text,title,dp) or not cover_from_pdf(pp,cp):return None
  return {'title':title,'category':cat,'source_page':url,'language':'id','source_group':group,'pdf_source':'/'+str(pp).replace('\\','/'),'docx_source':'/'+str(dp).replace('\\','/'),'cover_url':'/'+str(cp).replace('\\','/'),'type':'pdf'}
 except Exception:return None

def classify_pdf_title(title):
 t=title.lower()
 if KH.search(t): return 'khutbah'
 if KU.search(t): return 'kultum'
 if TF.search(t): return 'tafsir'
 return 'buku'

def rank(r):
 group={'muhammadiyah':10000,'kemenag':8000,'salafi':7000,'dewan-dakwah':6000,'nu':5000,'pesantren':4000,'islam':3000}.get(r.get('source_group','islam'),1000)
 lang=0 if str(r.get('language','')).lower() in ('id','indonesia','indonesian','bahasa indonesia') else 1
 y=-date_year(r.get('date'))
 return (lang,-group,y,clean_title(r.get('title','')).casefold())

def build():
 rows=[]
 # 1) Muhammadiyah article streams: newest official Indonesian material first.
 terms={'khutbah':['Khutbah Jumat','Khutbah Idul Fitri','Khutbah Gerhana'], 'kultum':['kultum','ceramah singkat','kuliah tujuh menit','pidato agama'], 'tafsir':['tafsir','Tafsir Al-Quran','Tafsir Al-Qur’an']}
 for section,ts in terms.items():
  for term in ts:
   for base in ['https://muhammadiyah.or.id']:
    for p in wp_posts(base,term,25):
     r=article_record(section,p.get('link',''),'muhammadiyah')
     if r: r['date']=p.get('date') or p.get('modified') or r.get('date',''); rows.append(r)
 # HTML category/search crawl fallback and extra older pages.
 for section in ('khutbah','kultum','tafsir'):
  for start in STARTS[section]:
   for typ,u,txt in discover_html(start,section,25):
    if typ!='pdf':
     r=article_record(section,u,'muhammadiyah')
     if r: rows.append(r)
    else:
     rr=remote_pdf(txt or Path(urlparse(u).path).stem.replace('-',' ').title(),u,section,'muhammadiyah')
     if rr: rows.append(rr)
 # 2) Direct Tarjih PDFs: mirror only real public PDFs and generate DOCX/cover locally.
 for section,title,url,group in KNOWN:
  rr=remote_pdf(title,url,section,group)
  if rr: rows.append(rr)
 # crawl Tarjih download area for additional real PDFs.
 for typ,u,txt in discover_html(STARTS['buku'][0],'buku',35):
  if typ!='pdf':continue
  title=clean_title(txt or Path(urlparse(u).path).stem.replace('-',' ').title()); cat=classify_pdf_title(title)
  if cat in ('khutbah','kultum','tafsir'): continue
  if not in_section(title,'buku'):continue
  rr=remote_pdf(title,u,'buku','muhammadiyah')
  if rr: rows.append(rr)
 # keep only complete local records; retain existing local assets.
 old=json.loads(OUT.read_text(encoding='utf-8')).get('items',[]) if OUT.exists() else []
 for r in old:
  if all(str(r.get(k,'')).startswith('/assets/') for k in ('pdf_source','docx_source','cover_url')): rows.append(r)
 seen={}
 for r in rows:
  title=clean_title(r.get('title')); k=re.sub(r'\W+','',title.casefold())
  if not title or k in seen: continue
  if not in_section(title,r.get('category','buku')): continue
  if not all(str(r.get(k2,'')).startswith('/assets/') for k2 in ('pdf_source','docx_source','cover_url')): continue
  seen[k]=r
 items=sorted(seen.values(),key=rank)[:MAX]
 OUT.write_text(json.dumps({'generated':time.strftime('%Y-%m-%dT%H:%M:%SZ',time.gmtime()),'total':len(items),'items':items},ensure_ascii=False,indent=2)+'\n',encoding='utf-8')
 print('usable PustakaMu records:',len(items))

if __name__=='__main__': build()
