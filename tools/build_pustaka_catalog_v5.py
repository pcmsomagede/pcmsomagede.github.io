#!/usr/bin/env python3
"""Build a real PustakaMu catalog.

Key rules:
- never reconstruct source PDFs when an original PDF exists;
- HTML articles are exported with a shaping-capable Arabic renderer;
- tafsir is a bibliography of tafsir books, not news articles;
- books/khutbah/kultum are gathered from several Islamic/Muhammadiyah sources;
- remote PDFs remain valid catalog entries when too large to cache locally.
"""
from __future__ import annotations
import hashlib, json, os, re, time
from pathlib import Path
from urllib.parse import urljoin, urlparse
import requests
from bs4 import BeautifulSoup
from docx import Document

OUT=Path('data/pustaka-catalog.json')
PDF=Path('assets/pustaka-files'); DOCX=Path('assets/pustaka-docx'); COV=Path('assets/pustaka-covers')
MAX_ITEMS=int(os.getenv('PUSTAKA_MAX_ITEMS','1000000'))
MAX_PDF_MB=int(os.getenv('PUSTAKA_LOCAL_MAX_MB','25'))
PER=100
S=requests.Session(); S.headers.update({'User-Agent':'PCM-Somagede-Pustaka/15.0 (+https://pcmsomagede.github.io/)'})
AR=re.compile(r'[\u0600-\u06ff\u0750-\u077f\u08a0-\u08ff]')
NOISE=re.compile(r'computer|programming|python|astronomy|physics|chemistry|mathematics|business|marketing|finance|engineering|medicine|biology|travel|cooking|recipe|novel|romance|crime|fiction',re.I)
KH=re.compile(r'khutbah|khutba|jumat|jum.?at|jumah|idul\s+(fitri|adha)|gerhana|istisqa|istisqo',re.I)
KU=re.compile(r'kultum|kuliah\s+tujuh\s+menit|7\s*menit|ceramah\s+(singkat|pendek)|pidato\s+agama',re.I)
TF=re.compile(r'\btafsir\b|tanzil|ta.?wil',re.I)
HD=re.compile(r'hadits?|hadith|sahih\s+bukhari|sahih\s+muslim|sunan\s+abu|tirmidzi|tirmidhi|nasai|nasa.?i|ibnu\s+majah|riyad|bulugh',re.I)
ISLAM=re.compile(r'islam|muslim|qur.?an|hadis|hadith|tafsir|fiqih|fikih|aqidah|akidah|tauhid|sunnah|syariah|dakwah|khutbah|jumat|kultum|ceramah|pengajian|pesantren|kitab|sholat|shalat|zakat|puasa|haji|umrah|ulama|ustadz|ustaz|ibadah|akhlak|tarjih|muhammadiyah|aisyiyah',re.I)
SOURCES=[
 'https://muhammadiyah.or.id','https://tarjih.or.id','https://buku.muhammadiyah.or.id',
 'https://khazanah.muhammadiyah.or.id','https://suaramuhammadiyah.id','https://www.dakwah.id'
]
TAFSIR_SEED=[
 ('Tafsir Ath-Thabari','Abu Ja’far Muhammad bin Jarir ath-Thabari','klasik',None,'Jami al-Bayan ‘an Ta’wil Ayi al-Quran'),
 ('Tafsir Al-Baghawi','Abu Muhammad al-Hasan bin Mas’ud al-Baghawi','klasik',None,'Ma’alim at-Tanzil'),
 ('Tafsir Al-Muharrar','Abu Muhammad Abdul Haq bin Ghalib bin ‘Athiyah','klasik',None,'Al-Muharrar al-Wajiz fi Tafsir al-Kitab al-‘Aziz'),
 ('Tafsir Ibnu Katsir','Abu al-Fida’ Ismail bin Umar Ibnu Katsir','klasik',None,'Tafsir al-Quran al-Azhim'),
 ('Tafsir Ad-Durr Al-Mantsur','Jalaluddin as-Suyuthi','klasik',None,'Ad-Durr al-Matsur fi at-Tafsir bi al-Ma’tsur'),
 ('Tafsir Fathul Qadir','Muhammad bin Ali asy-Syaukani','klasik',None,'Fathul Qadir'),
 ('Tafsir Adhwa’ Al-Bayan','Muhammad al-Amin bin Muhammad al-Mukhtar asy-Syinqithi','klasik',None,'Adhwa’ al-Bayan fi Idhah al-Quran bi al-Quran'),
 ('Tafsir Ar-Razi','Fakhruddin ar-Razi','klasik',None,'Mafatih al-Ghaib'),
 ('Tafsir Al-Baidhawi','Abdullah bin Umar al-Baidhawi','klasik',None,'Anwar at-Tanzil wa Asrar at-Ta’wil'),
 ('Tafsir An-Nasafi','Abu al-Barakat Abdullah bin Ahmad an-Nasafi','klasik',None,'Madarik at-Tanzil wa Haqaiq at-Ta’wil'),
 ('Tafsir Al-Khazin','Alauddin Ali bin Muhammad al-Khazin','klasik',None,'Lubab at-Ta’wil fi Ma’ani at-Tanzil'),
 ('Tafsir Bahrul Muhith','Muhammad bin Yusuf Abu Hayyan','klasik',None,'Bahrul Muhith'),
 ('Tafsir Al-Qurthubi','Abu Abdullah Muhammad bin Ahmad al-Anshari al-Qurthubi','klasik',None,'Al-Jami’ li Ahkam al-Quran'),
 ('Tafsir Jalalain','Jalaluddin al-Mahalli dan Jalaluddin as-Suyuthi','klasik','870 H','Tafsir al-Jalalain'),
 ('Tafsir Al-Muyassar','Tim ulama Markaz Tafsir','kontemporer',None,'Tafsir al-Muyassar'),
 ('Al-Mukhtashar fi at-Tafsir','Tim ulama Markaz Tafsir Liddirasat al-Qur’aniyah','kontemporer',None,'Al-Mukhtashar fi at-Tafsir'),
 ('Tafsir As-Sa’di','Abdurrahman bin Nashir as-Sa’di','klasik',None,'Taisir al-Karim ar-Rahman fi Tafsiri Kalami al-Mannan'),
 ('Turjuman al-Mustafid','Abdur Rauf al-Fanshuri as-Singkili','Nusantara','abad ke-17 M','Turjuman al-Mustafid'),
 ('Marah Labid / At-Tafsir al-Munir','Muhammad Nawawi al-Bantani al-Jawi','Nusantara','1305 H','Marah Labid li Kasyfi Ma’na Quran Majid'),
 ('Faidhu ar-Rahman','Muhammad Shalih Darat as-Samarani','Nusantara',None,'Faidhu ar-Rahman'),
 ('Tafsir Al-Furqan','Ahmad Hasan Bandung','Nusantara','1928 M','Tafsir Al-Furqan'),
 ('Tafsir Al-Quran','Mahmud Yunus','Nusantara','1935 M','Tafsir al-Quran'),
 ('Tafsir An-Nuur','Hasbi ash-Shidqi','Nusantara','1956 M','Tafsir an-Nuur'),
 ('Tafsir Al-Bayan','Hasbi ash-Shidqi','Nusantara','1956 M','Tafsir al-Bayan'),
 ('Tafsir Al-Quran al-Karim','Halim Hasan','Nusantara','1955 M','Tafsir al-Quran al-Karim'),
 ('Al-Ibriz li Ma’rifati al-Quran','Bishri Musthafa','Nusantara','1960 M','Al-Ibriz li Ma’rifati al-Quran'),
 ('Tafsir Raudhatul Irfan','Ahmad Sanusi', 'Nusantara', None,'Tafsir Raudhatul Irfan fi Ma’rifati al-Quran'),
 ('Tafsir Al-Azhar','Abdul Malik Abdul Karim Amrullah (HAMKA)','Nusantara',None,'Tafsir Al-Azhar'),
 ('Ahkam al-Quran','Al-Jashash','tafsir ayat ahkam',None,'Ahkam al-Quran'),
 ('Ahkam al-Quran','Ibnul Arabi','tafsir ayat ahkam',None,'Ahkam al-Quran'),
 ('Al-Iklil fi Istinbath at-Tanzil','Jalaluddin as-Suyuthi','tafsir ayat ahkam',None,'Al-Iklil fi Istinbath at-Tanzil'),
 ('Tafsir Ayat al-Ahkam','As-Sayus','tafsir ayat ahkam',None,'Tafsir Ayat al-Ahkam'),
 ('Tafsir Ayat al-Ahkam','Manna’ al-Qathan','tafsir ayat ahkam',None,'Tafsir Ayat al-Ahkam'),
 ('Rawa’i’ al-Bayan Tafsir Ayat al-Ahkam','Muhammad Ali ash-Shabuni','tafsir ayat ahkam',None,'Rawa’i’ al-Bayan Tafsir Ayat al-Ahkam'),
 ('At-Tafsir wa al-Bayan','Abdul Aziz bin Marzuq ath-Tharifi','tafsir ayat ahkam',None,'At-Tafsir wa al-Bayan'),
]

def slug(t):
 s=re.sub(r'[^a-z0-9]+','-',t.lower()).strip('-'); return (s[:110] or hashlib.sha1(t.encode()).hexdigest()[:16])
def clean(t): return re.sub(r'\s+',' ',str(t or '')).strip(' -|')
def local(v):
 s=str(v or ''); return bool(s) and (s.startswith('/') and Path(s.lstrip('/')).exists())
def remote(v): return bool(re.match(r'^https?://',str(v or '')))

def sanitize_html(html,url,title):
 s=BeautifulSoup(html,'html.parser')
 for z in s.select('script,style,noscript,nav,header,footer,form,svg,iframe,.share,.sharing,.comments,.comment,.related,.advertisement,.ads,.sidebar'): z.decompose()
 root=s.select_one('article') or s.select_one('.entry-content') or s.select_one('.post-content') or s.select_one('main') or s.body or s
 keep=[]
 for e in root.find_all(['h1','h2','h3','h4','p','ol','ul','li','blockquote','strong','em','br','table','tr','td','th']):
  txt=e.get_text(' ',strip=True)
  if txt and (len(txt)>1):
   if e.name in {'p','li','blockquote'} and AR.search(txt): e['dir']='rtl'
   keep.append(e)
 body=''.join(str(x) for x in keep) or f'<h1>{title}</h1><p>{clean(root.get_text(" ",strip=True))}</p>'
 return f'''<!doctype html><html lang="id"><head><meta charset="utf-8"><style>@page{{size:A4;margin:17mm 16mm}}body{{font-family:"Noto Sans","Noto Sans Arabic","Noto Naskh Arabic",sans-serif;font-size:11pt;line-height:1.65;color:#171717}}h1,h2,h3{{font-family:"Noto Sans",sans-serif;line-height:1.25}}p,li{{margin:.45em 0}}[dir=rtl]{{direction:rtl;unicode-bidi:embed;text-align:right;font-family:"Noto Naskh Arabic","Noto Sans Arabic",serif;font-size:14pt;line-height:2}}blockquote{{border-left:3px solid #aaa;padding-left:10px}}table{{border-collapse:collapse;width:100%}}td,th{{border:1px solid #bbb;padding:4px}}</style></head><body><h1>{clean(title)}</h1>{body}</body></html>'''

def html_export(title,url):
 try:
  from weasyprint import HTML
  r=S.get(url,timeout=60); r.raise_for_status(); data=sanitize_html(r.text,url,title)
  n=slug(title); pp=PDF/(n+'.pdf'); dp=DOCX/(n+'.docx'); cp=COV/(n+'.jpg'); pp.parent.mkdir(parents=True,exist_ok=True)
  HTML(string=data,base_url=url).write_pdf(pp)
  soup=BeautifulSoup(data,'html.parser'); d=Document(); d.add_heading(title,0)
  for e in soup.find_all(['p','li','blockquote']):
   t=e.get_text(' ',strip=False)
   if t.strip():
    p=d.add_paragraph(t)
    if AR.search(t): p.alignment=2
  d.save(dp)
  try:
   import fitz
   doc=fitz.open(pp); pix=doc[0].get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False); cp.parent.mkdir(parents=True,exist_ok=True); pix.save(cp); doc.close()
  except Exception: pass
  if not pp.exists() or pp.stat().st_size<2000:return None
  return {'title':title,'category':'','source_page':url,'language':'id','pdf_source':'/'+str(pp).replace('\\','/'),'docx_source':'/'+str(dp).replace('\\','/'),'cover_url':'/'+str(cp).replace('\\','/') if cp.exists() else '','type':'article-export'}
 except Exception as e:
  print('HTML export failed',url,e); return None

def get_json(url,params=None):
 try:
  r=S.get(url,params=params,timeout=45); r.raise_for_status(); return r.json()
 except Exception:return None

def wp_posts(base,terms,max_pages=60):
 out=[]; seen=set(); api=base.rstrip('/')+'/wp-json/wp/v2/posts'
 for term in terms:
  for p in range(1,max_pages+1):
   rows=get_json(api,{'search':term,'per_page':PER,'page':p,'_fields':'link,date,modified,title,content,featured_media'})
   if not rows: break
   for row in rows:
    u=row.get('link','')
    if u and u not in seen:seen.add(u);out.append(row)
   if len(rows)<PER:break
 return out

def pdf_links(base,max_pages=40):
 urls=[];seen=set(); queue=[base.rstrip('/')]
 for _ in range(max_pages):
  if not queue:break
  u=queue.pop(0)
  if u in seen:continue
  seen.add(u)
  try:
   r=S.get(u,timeout=45);r.raise_for_status();s=BeautifulSoup(r.text,'html.parser')
   for a in s.select('a[href]'):
    h=urljoin(u,a.get('href','')); text=clean(a.get_text(' ',strip=True))
    if urlparse(h).netloc not in urlparse(base).netloc and urlparse(h).netloc!=urlparse(base).netloc:continue
    if re.search(r'\.pdf(?:$|\?)',h,re.I):
     if h not in seen and h not in {x[0] for x in urls}:urls.append((h,text or Path(urlparse(h).path).stem))
    elif any(k in (text+' '+h).lower() for k in ['buku','kitab','tarjih','download','pustaka','khutbah','kultum']):
     if h not in seen and len(queue)<250:queue.append(h)
  except Exception:continue
 return urls

def pdf_record(title,url,cat,group):
 n=slug(title); pp=PDF/(n+'.pdf'); dp=DOCX/(n+'.docx'); cp=COV/(n+'.jpg')
 try:
  rr=S.get(url,stream=True,timeout=90); rr.raise_for_status(); size=int(rr.headers.get('content-length') or 0)
  if size and size<=MAX_PDF_MB*1024*1024:
   data=rr.content; pp.parent.mkdir(parents=True,exist_ok=True);pp.write_bytes(data)
   try:
    import fitz
    doc=fitz.open(pp); txt='\n\n'.join(pg.get_text('text') for pg in doc); doc.save(pp.with_suffix('.tmp.pdf')); doc.close()
    # keep original bytes, only use text to build DOCX
    d=Document();d.add_heading(title,0)
    for para in re.split(r'\n\s*\n',txt):
     if para.strip(): d.add_paragraph(para.strip())
    dp.parent.mkdir(parents=True,exist_ok=True);d.save(dp)
    doc=fitz.open(pp);pix=doc[0].get_pixmap(matrix=fitz.Matrix(1.25,1.25),alpha=False);cp.parent.mkdir(parents=True,exist_ok=True);pix.save(cp);doc.close()
   except Exception: pass
   return {'title':title,'category':cat,'source_page':url,'language':'id','source_group':group,'pdf_source':'/'+str(pp).replace('\\','/'),'docx_source':'/'+str(dp).replace('\\','/') if dp.exists() else '','cover_url':'/'+str(cp).replace('\\','/') if cp.exists() else '','type':'pdf-original'}
  rr.close()
  return {'title':title,'category':cat,'source_page':url,'language':'id','source_group':group,'pdf_source':url,'docx_source':'','cover_url':'','type':'pdf-remote'}
 except Exception:return None

def seed_tafsir():
 out=[]
 src='https://www.dakwah.id/kumpulan-kitab-tafsir-terpopuler-klasik-dan-kontemporer/'
 for title,author,period,year,full in TAFSIR_SEED:
  out.append({'title':title,'author':author,'year':year or '', 'period':period,'full_title':full,'category':'tafsir','source_page':src,'source_group':'dakwah.id','language':'id','type':'tafsir-book','status':'metadata'})
 return out

def discover_articles():
 out=[]
 terms={'khutbah':['khutbah jumat','khutbah idul fitri','khutbah idul adha','khutbah gerhana'], 'kultum':['kultum','ceramah singkat','kuliah tujuh menit','pidato agama'], 'buku':['himpunan putusan tarjih','tanya jawab agama','manhaj tarjih','tuntunan','pedoman','buku islam','kitab'],}
 for sec,ts in terms.items():
  for base in ['https://muhammadiyah.or.id','https://tarjih.or.id']:
   for row in wp_posts(base,ts,45):
    title=clean((row.get('title') or {}).get('rendered') if isinstance(row.get('title'),dict) else row.get('title'))
    if not title or (NOISE.search(title) and not ISLAM.search(title)):continue
    t=KH.search(title) if sec=='khutbah' else KU.search(title) if sec=='kultum' else None
    if sec!='buku' and not t:continue
    if sec=='buku' and (KH.search(title) or KU.search(title) or TF.search(title) or HD.search(title)):continue
    r=html_export(title,row.get('link',''))
    if r:r['category']=sec;r['source_group']='muhammadiyah' if 'muhammadiyah.or.id' in row.get('link','') else 'tarjih';r['date']=row.get('date') or row.get('modified') or '';out.append(r)
 # Dedicated PDF/flipbook discovery from Muhammadiyah ecosystem.
 for base in ['https://tarjih.or.id','https://buku.muhammadiyah.or.id','https://khazanah.muhammadiyah.or.id']:
  cat='buku'
  for u,lab in pdf_links(base,35):
   title=clean(lab)
   if not title or NOISE.search(title) or KH.search(title) or KU.search(title) or TF.search(title) or HD.search(title):continue
   if not ISLAM.search(title) and not re.search(r'tarjih|himpunan|tuntunan|pedoman|muhammadiyah|kitab|buku',title,re.I):continue
   r=pdf_record(title,u,cat,'muhammadiyah');
   if r:out.append(r)
 return out

def build():
 rows=seed_tafsir()+discover_articles()
 # Preserve any already-built valid entries that are not duplicated by title/category.
 if OUT.exists():
  try: rows += json.loads(OUT.read_text(encoding='utf-8')).get('items',[])
  except Exception: pass
 seen=set(); final=[]
 for x in rows:
  x['title']=clean(x.get('title') or x.get('name'))
  if not x['title']:continue
  key=(x['category'] or x.get('type',''),x['title'].casefold())
  if key in seen:continue
  if len(final)>=MAX_ITEMS:break
  seen.add(key);final.append(x)
 final.sort(key=lambda x:(str(x.get('category','')),str(x.get('year') or ''),x['title'].casefold()))
 OUT.parent.mkdir(parents=True,exist_ok=True); OUT.write_text(json.dumps({'generated':time.strftime('%Y-%m-%dT%H:%M:%SZ',time.gmtime()),'total':len(final),'items':final},ensure_ascii=False,indent=2)+'\n',encoding='utf-8')
 print('PustakaMu candidate records:',len(final))
if __name__=='__main__':build()
